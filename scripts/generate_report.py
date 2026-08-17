#!/usr/bin/env python3
"""generate_report.py — Parameterized UCITS Outlook email report generator.

Reads three input files (all user-supplied, NO hardcoded fund/portfolio/login data):

  config.json     fund name, title, compliance thresholds, fonts, language, paths
  events.json     the daily analysis content (per-run input — the "brains")
  holdings.csv    portfolio positions

Produces a 5-part .docx:
  Cover -> Part 1 Major Events -> Part 2 Investment Summary ->
  Part 3 Long Ideas -> Part 4 Portfolio Table + UCITS compliance check ->
  Part 5 Link Status.

Usage:
  python generate_report.py [path/to/config.json]

All paths in config.json are resolved relative to the config file's directory,
so the skill folder is portable.
"""
import json, os, csv, sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def load_config(path):
    with open(path, "r", encoding="utf-8") as f:
        cfg = json.load(f)
    base = os.path.dirname(os.path.abspath(path))
    for key in ("output_docx", "events_json", "holdings_csv"):
        p = cfg["paths"][key]
        if not os.path.isabs(p):
            cfg["paths"][key] = os.path.join(base, p)
    return cfg


def rgb(hex_str):
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


class ReportBuilder:
    def __init__(self, cfg):
        self.cfg = cfg
        self.r = cfg["report"]
        self.deep = rgb(self.r.get("deep_blue_hex", "003D6B"))
        self.med = rgb(self.r.get("medium_blue_hex", "005B8E"))
        self.light = rgb(self.r.get("light_blue_hex", "D6E8F7"))
        self.font = self.r.get("font_latin", "Calibri")
        self.font_east = self.r.get("font_eastasia", "Microsoft YaHei")
        self.labels = cfg.get("labels", {})
        self.doc = Document()
        for s in self.doc.sections:
            s.top_margin = Cm(2.0)
            s.bottom_margin = Cm(2.0)
            s.left_margin = Cm(2.5)
            s.right_margin = Cm(2.5)
        self.ev = {}
        self.holdings = []

    # ---- low-level helpers ----
    def para(self, text, size=11, bold=False, color=None, align=None,
             space_after=6, space_before=0, font=None):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        if align is not None:
            p.alignment = align
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = font or self.font
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_east)
        run.bold = bold
        if color:
            run.font.color.rgb = color
        return p

    def section(self, title, size=18):
        self.para(title, size=size, bold=True, color=self.deep, space_before=18, space_after=8)

    def sub(self, title, size=13):
        self.para(title, size=size, bold=True, color=self.med, space_before=12, space_after=4)

    def body(self, text, size=10, space_after=4):
        self.para(text, size=size, space_after=space_after)

    def event_block(self, ev):
        imp = ev.get("importance", self.labels.get("importance_red", "🔴"))
        self.para(f'{imp} {ev.get("title", "")}', size=11, bold=True,
                  color=self.deep, space_before=10, space_after=2)
        self.body(f'⏱ {self.labels.get("time_source", "Time / Source")}: {ev.get("time_source", "")}')
        self.body(f'📊 {self.labels.get("key_data", "Key Data")}: {ev.get("key_data", "")}')
        self.body(f'📝 {self.labels.get("analysis", "Analysis")}: {ev.get("analysis", "")}')

    # ---- sections ----
    def cover(self):
        for _ in range(4):
            self.para("", size=11, space_after=0)
        self.para(self.r["title"], size=28, bold=True, color=self.deep, align=WD_ALIGN_PARAGRAPH.CENTER)
        self.para(self.r.get("subtitle", ""), size=14, color=self.deep,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
        self.para(f'{self.r.get("report_date", "<DATE>")} | {self.r.get("manager", "<FIRM>")}',
                  size=14, color=self.deep, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
        self.para(self.r.get("confidential_note", "CONFIDENTIAL — For Internal Use Only"),
                  size=11, color=self.deep, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
        self.para("", size=11)
        cs = self.ev.get("cover_stats", {})
        self.para(f'Email sources: {cs.get("sources_line", "<your broker sources>")}',
                  size=9, color=self.deep, align=WD_ALIGN_PARAGRAPH.CENTER)
        self.para(f'Total fetched: {cs.get("total_emails", "?")} | Deduped: {cs.get("deduped", "?")} | '
                  f'GC-relevant: {cs.get("gc_relevant", "?")} | Non-GC: {cs.get("non_gc", "?")}',
                  size=9, color=self.deep, align=WD_ALIGN_PARAGRAPH.CENTER)
        self.para("nolinks fetch — report_links_crawled: false", size=9, color=self.deep,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        self.doc.add_page_break()

    def part1(self):
        self.section(self.labels.get("part1", "Part 1   Major Market Events"))
        cs = self.ev.get("cover_stats", {})
        self.body(f'Report date: {self.r.get("report_date", "<DATE>")} | '
                  f'Emails covered: {cs.get("coverage_date", "<DATE>")} | '
                  f'GC-relevant emails: ~{cs.get("gc_relevant", "?")}', space_after=8)
        for sec in self.ev.get("part1_sections", []):
            self.sub(sec.get("subheading", ""))
            for ev in sec.get("events", []):
                self.event_block(ev)

    def part2(self):
        self.doc.add_page_break()
        self.section(self.labels.get("part2", "Part 2   Investment Ideas Summary"))
        self.body(self.ev.get("part2_summary", ""), space_after=8)

    def part3(self):
        self.section(self.labels.get("part3", "Part 3   Long Ideas"))
        for li in self.ev.get("part3_long_ideas", []):
            self.sub(li.get("name", ""))
            self.body(li.get("subtitle", ""), space_after=2)
            self.body(li.get("body", ""), space_after=6)

    def part4(self):
        self.doc.add_page_break()
        self.section(self.labels.get("part4", "Part 4   Portfolio Allocation"))
        self.body(self.ev.get("part4_note", "Based on latest available holdings."), space_after=8)
        self.portfolio_table()
        self.compliance_check()

    def portfolio_table(self):
        headers = ["Name / Ticker", "Latest", "Entry", "Target", "Weight %", "Compliant"]
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            for p in hdr[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(8)
                    run.font.name = self.font
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_east)
            set_cell_shading(hdr[i], self.r.get("deep_blue_hex", "003D6B"))
            for p in hdr[i].paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.bold = True
        for row in self.holdings:
            cells = table.add_row().cells
            vals = [row.get("name", ""), row.get("latest", ""), row.get("entry", ""),
                    row.get("target", ""), row.get("weight", ""), row.get("compliance", "✅")]
            for i, v in enumerate(vals):
                cells[i].text = str(v)
                for p in cells[i].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(8)
                        run.font.name = self.font
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_east)
                if str(row.get("compliance", "")) not in ("✅", "", "PASS"):
                    set_cell_shading(cells[i], "FFF2CC")

    def compliance_check(self):
        self.body("", size=6)
        self.sub("UCITS Compliance Check", size=12)
        u = self.cfg.get("ucits", {})
        rows = self.holdings
        cap = float(u.get("single_name_cap_pct", 10))
        eq = [r for r in rows if str(r.get("region", "")).lower() != "cash"]
        n = len(eq)
        mn = int(u.get("stock_count_min", 23))
        mx = int(u.get("stock_count_max", 27))
        gc_regions = {"hk", "a-share", "taiwan", "cash"}
        non_gc = [r for r in rows if str(r.get("region", "")).lower() not in gc_regions]
        over = [r for r in eq if float(r.get("weight", 0) or 0) > cap]
        big = [r for r in eq if float(r.get("weight", 0) or 0) > 5]
        s540 = sum(float(r.get("weight", 0) or 0) for r in big)
        ash = sum(float(r.get("weight", 0) or 0) for r in rows if str(r.get("region", "")).lower() == "a-share")
        tw = sum(float(r.get("weight", 0) or 0) for r in rows if str(r.get("region", "")).lower() == "taiwan")
        cash = sum(float(r.get("weight", 0) or 0) for r in rows if str(r.get("region", "")).lower() == "cash")
        flagged = [r for r in rows if r.get("name", "") in (u.get("flagged_names", []) or [])]
        ok = "✅"; no = "❌"
        lines = []
        lines.append(f'{"✅" if mn <= n <= mx else no} Stock count: {n} (rule {mn}-{mx})')
        if u.get("greater_china_only") and non_gc:
            names = ", ".join(r.get("name", "?") for r in non_gc)
            lines.append(f'{no} Eligible-region only: {names} — VIOLATION (non-eligible exposure)')
        else:
            lines.append(f'{ok} Eligible-region only: all positions in eligible regions')
        maxw = max([float(r.get("weight", 0) or 0) for r in eq] + [0])
        lines.append(f'{"✅" if not over else no} Single-name ≤{cap}%: max {maxw:.2f}%')
        if u.get("rule_5_40"):
            lines.append(f'{"✅" if s540 <= 40 else no} 5/40 rule: top>5% sum = {s540:.2f}% (≤40%)')
        if u.get("a_share_cap_pct") is not None:
            lines.append(f'{"✅" if ash < float(u["a_share_cap_pct"]) else no} A-shares: {ash:.2f}% (≤{u["a_share_cap_pct"]}%)')
        if u.get("taiwan_cap_pct") is not None:
            lines.append(f'{"✅" if tw < float(u["taiwan_cap_pct"]) else no} Taiwan: {tw:.2f}% (≤{u["taiwan_cap_pct"]}%)')
        if u.get("cash_cap_pct") is not None:
            lines.append(f'{"✅" if cash < float(u["cash_cap_pct"]) else no} Cash: {cash:.2f}% (≤{u["cash_cap_pct"]}%)')
        if flagged:
            lines.append(f'{no} Flagged names: {", ".join(r.get("name", "?") for r in flagged)} — {u.get("flagged_note", "")}')
        self.body("\n".join(lines), space_after=8)

    def part5(self):
        self.doc.add_page_break()
        self.section(self.labels.get("part5", "Part 5   Link Status"))
        self.body(self.ev.get("part5_link_status", ""), space_after=8)
        self.para("— END OF REPORT —", size=10, color=self.deep,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_before=20)
        self.para(f'Generated by report skill | {self.r.get("manager", "<FIRM>")} | '
                  f'{self.r.get("confidential_note", "CONFIDENTIAL")}',
                  size=8, color=self.deep, align=WD_ALIGN_PARAGRAPH.CENTER)

    def load_holdings(self, path):
        rows = []
        with open(path, "r", encoding="utf-8-sig") as f:
            for row in csv.DictReader(f):
                rows.append(row)
        return rows

    def build(self):
        with open(self.cfg["paths"]["events_json"], "r", encoding="utf-8") as f:
            self.ev = json.load(f)
        self.holdings = self.load_holdings(self.cfg["paths"]["holdings_csv"])
        self.cover()
        self.part1()
        self.part2()
        self.part3()
        self.part4()
        self.part5()
        out = self.cfg["paths"]["output_docx"]
        os.makedirs(os.path.dirname(out) or ".", exist_ok=True)
        self.doc.save(out)
        print(f"✅ Report saved: {out}")


def main():
    cfg_path = sys.argv[1] if len(sys.argv) > 1 else "config.example.json"
    cfg = load_config(cfg_path)
    ReportBuilder(cfg).build()


if __name__ == "__main__":
    main()
